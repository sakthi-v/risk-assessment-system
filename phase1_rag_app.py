import streamlit as st
import google.generativeai as genai
from pathlib import Path
import tempfile
import os
import time
from typing import List, Dict
import json
import pickle

# Document processing libraries
from docx import Document
import PyPDF2
import pandas as pd
import openpyxl
from openpyxl import load_workbook

# Vector store and embeddings
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Configure page
st.set_page_config(
    page_title="RAG Document Q&A System",
    page_icon="📚",
    layout="wide"
)

# ===================================================================
# PERSISTENT STORAGE PATHS
# ===================================================================
KNOWLEDGE_BASE_DIR = Path("knowledge_base")
KNOWLEDGE_BASE_DIR.mkdir(exist_ok=True)

KB_DOCUMENTS_FILE = KNOWLEDGE_BASE_DIR / "documents.pkl"
KB_VECTORIZER_FILE = KNOWLEDGE_BASE_DIR / "vectorizer.pkl"
KB_VECTORS_FILE = KNOWLEDGE_BASE_DIR / "document_vectors.pkl"
KB_METADATA_FILE = KNOWLEDGE_BASE_DIR / "metadata.json"

# Initialize session state
if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'processed' not in st.session_state:
    st.session_state.processed = False
if 'vectorizer' not in st.session_state:
    st.session_state.vectorizer = None
if 'document_vectors' not in st.session_state:
    st.session_state.document_vectors = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'kb_loaded' not in st.session_state:
    st.session_state.kb_loaded = False

# ===================================================================
# KNOWLEDGE BASE PERSISTENCE FUNCTIONS
# ===================================================================

def save_knowledge_base(documents: List[Dict], vectorizer, document_vectors):
    """Save the entire knowledge base to disk"""
    try:
        # Save documents
        with open(KB_DOCUMENTS_FILE, 'wb') as f:
            pickle.dump(documents, f)
        
        # Save vectorizer
        with open(KB_VECTORIZER_FILE, 'wb') as f:
            pickle.dump(vectorizer, f)
        
        # Save document vectors
        with open(KB_VECTORS_FILE, 'wb') as f:
            pickle.dump(document_vectors, f)
        
        # Save metadata
        metadata = {
            'num_documents': len(documents),
            'document_names': [doc['filename'] for doc in documents],
            'saved_at': time.strftime('%Y-%m-%d %H:%M:%S'),
            'version': '1.0'
        }
        with open(KB_METADATA_FILE, 'w') as f:
            json.dump(metadata, f, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Error saving knowledge base: {str(e)}")
        return False

def load_knowledge_base():
    """Load the knowledge base from disk"""
    try:
        # Check if all files exist
        if not all([
            KB_DOCUMENTS_FILE.exists(),
            KB_VECTORIZER_FILE.exists(),
            KB_VECTORS_FILE.exists(),
            KB_METADATA_FILE.exists()
        ]):
            return None, None, None, None
        
        # Load documents
        with open(KB_DOCUMENTS_FILE, 'rb') as f:
            documents = pickle.load(f)
        
        # Load vectorizer
        with open(KB_VECTORIZER_FILE, 'rb') as f:
            vectorizer = pickle.load(f)
        
        # Load document vectors
        with open(KB_VECTORS_FILE, 'rb') as f:
            document_vectors = pickle.load(f)
        
        # Load metadata
        with open(KB_METADATA_FILE, 'r') as f:
            metadata = json.load(f)
        
        return documents, vectorizer, document_vectors, metadata
    except Exception as e:
        st.error(f"Error loading knowledge base: {str(e)}")
        return None, None, None, None

def knowledge_base_exists():
    """Check if a saved knowledge base exists"""
    return all([
        KB_DOCUMENTS_FILE.exists(),
        KB_VECTORIZER_FILE.exists(),
        KB_VECTORS_FILE.exists(),
        KB_METADATA_FILE.exists()
    ])

def delete_knowledge_base():
    """Delete the saved knowledge base"""
    try:
        for file in [KB_DOCUMENTS_FILE, KB_VECTORIZER_FILE, KB_VECTORS_FILE, KB_METADATA_FILE]:
            if file.exists():
                file.unlink()
        return True
    except Exception as e:
        st.error(f"Error deleting knowledge base: {str(e)}")
        return False

# ===================================================================
# DOCUMENT PROCESSOR (Same as before)
# ===================================================================

class DocumentProcessor:
    """Process different document types and extract text"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str, original_filename: str = None) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            display_name = original_filename if original_filename else Path(file_path).name
            
            text += f"\n{'='*80}\n"
            text += f"DOCUMENT: {display_name}\n"
            text += f"{'='*80}\n\n"
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text += f"Total Pages: {len(pdf_reader.pages)}\n\n"
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str, original_filename: str = None) -> str:
        """Extract text from DOCX file including tables with robust method"""
        text = ""
        try:
            display_name = original_filename if original_filename else Path(file_path).name
            
            doc = Document(file_path)
            
            text += f"\n{'='*80}\n"
            text += f"DOCUMENT: {display_name}\n"
            text += f"{'='*80}\n"
            text += f"Total Paragraphs: {len(doc.paragraphs)}\n"
            text += f"Total Tables: {len(doc.tables)}\n\n"
            
            text += f"{'='*80}\n"
            text += f"TEXT CONTENT\n"
            text += f"{'='*80}\n\n"
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            if len(doc.tables) > 0:
                text += f"\n{'='*80}\n"
                text += f"TABLES IN DOCUMENT ({len(doc.tables)} tables found)\n"
                text += f"{'='*80}\n\n"
                
                for table_idx, table in enumerate(doc.tables):
                    text += f"\n{'─'*80}\n"
                    text += f"TABLE {table_idx + 1}\n"
                    text += f"{'─'*80}\n"
                    
                    num_rows = len(table.rows)
                    num_cols = len(table.columns) if table.rows else 0
                    text += f"Dimensions: {num_rows} rows × {num_cols} columns\n\n"
                    
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace('\n', ' ')
                            row_data.append(cell_text if cell_text else "")
                        if any(row_data):
                            table_data.append(row_data)
                    
                    if table_data:
                        num_columns = len(table_data[0]) if table_data else 0
                        col_widths = [0] * num_columns
                        
                        for row in table_data:
                            for i, cell in enumerate(row):
                                if i < num_columns:
                                    col_widths[i] = max(col_widths[i], len(str(cell)))
                        
                        for row_idx, row in enumerate(table_data):
                            formatted_cells = []
                            for i, cell in enumerate(row):
                                if i < len(col_widths):
                                    formatted_cells.append(str(cell).ljust(col_widths[i]))
                                else:
                                    formatted_cells.append(str(cell))
                            
                            text += " | ".join(formatted_cells) + "\n"
                            
                            if row_idx == 0 and len(table_data) > 1:
                                separator_parts = []
                                for width in col_widths:
                                    separator_parts.append("─" * width)
                                text += "─┼─".join(separator_parts) + "\n"
                    else:
                        text += "(Empty table)\n"
                    
                    text += f"{'─'*80}\n\n"
            
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            text = f"[Error reading DOCX file: {str(e)}]"
        
        return text
    
    @staticmethod
    def extract_text_from_excel(file_path: str, original_filename: str = None) -> str:
        """Extract text from Excel file with formulas and detailed information"""
        text = ""
        try:
            display_name = original_filename if original_filename else Path(file_path).name
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.xlsx':
                try:
                    wb = load_workbook(file_path, data_only=False)
                except Exception as wb_error:
                    st.warning(f"Could not load formulas, using values only: {str(wb_error)}")
                    return DocumentProcessor._extract_excel_with_pandas(file_path, display_name)
                
                text += f"\n{'='*80}\n"
                text += f"DOCUMENT: {display_name}\n"
                text += f"{'='*80}\n"
                text += f"File Type: .xlsx (Excel 2007+)\n"
                text += f"Total Sheets: {len(wb.sheetnames)}\n"
                text += f"Sheet Names: {', '.join(wb.sheetnames)}\n\n"
                
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    
                    text += f"\n{'='*80}\n"
                    text += f"SHEET: {sheet_name}\n"
                    text += f"{'='*80}\n\n"
                    
                    max_row = sheet.max_row
                    max_col = sheet.max_column
                    text += f"Sheet Dimensions: {max_row} rows × {max_col} columns\n"
                    text += f"Data Range: A1 to {sheet.cell(max_row, max_col).coordinate}\n\n"
                    
                    formulas_in_sheet = []
                    text += "FORMULAS FOUND IN THIS SHEET:\n"
                    text += "-" * 80 + "\n"
                    
                    formula_count = 0
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                                formula_count += 1
                                formula = cell.value
                                try:
                                    wb_values = load_workbook(file_path, data_only=True)
                                    calculated_value = wb_values[sheet_name][cell.coordinate].value
                                    wb_values.close()
                                except:
                                    calculated_value = "N/A"
                                
                                formula_info = f"Cell {cell.coordinate}: {formula}"
                                if calculated_value != "N/A":
                                    formula_info += f" → Result: {calculated_value}"
                                
                                formulas_in_sheet.append(formula_info)
                                text += f"  {formula_info}\n"
                    
                    if formula_count == 0:
                        text += "  (No formulas found in this sheet)\n"
                    else:
                        text += f"\n  Total Formulas: {formula_count}\n"
                    
                    text += "\n"
                    
                    text += "COMPLETE SHEET DATA:\n"
                    text += "-" * 80 + "\n"
                    
                    for row_idx in range(1, max_row + 1):
                        row_cells = []
                        has_content = False
                        
                        for col_idx in range(1, max_col + 1):
                            cell = sheet.cell(row_idx, col_idx)
                            cell_value = cell.value
                            
                            if cell_value is not None:
                                has_content = True
                                if isinstance(cell_value, str) and cell_value.startswith('='):
                                    row_cells.append(f"[{cell_value}]")
                                else:
                                    row_cells.append(str(cell_value))
                            else:
                                row_cells.append("")
                        
                        if has_content:
                            text += " | ".join(row_cells) + "\n"
                    
                    if max_row > 500:
                        text += f"\n... (showing first 500 of {max_row} rows)\n"
                    
                    text += "\n"
                    
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        
                        if not df.empty and not df.dropna(how='all').empty:
                            text += "COLUMN INFORMATION:\n"
                            text += "-" * 80 + "\n"
                            text += f"Columns: {list(df.columns)}\n"
                            text += f"Total Rows with Data: {len(df)}\n\n"
                            
                            numeric_cols = df.select_dtypes(include=['number']).columns
                            if len(numeric_cols) > 0:
                                text += "NUMERIC COLUMN STATISTICS:\n"
                                text += "-" * 80 + "\n"
                                for col in numeric_cols:
                                    values = df[col].dropna()
                                    if len(values) > 0:
                                        text += f"\n{col}:\n"
                                        text += f"  Count: {len(values)}\n"
                                        text += f"  Sum: {values.sum():.2f}\n"
                                        text += f"  Mean: {values.mean():.2f}\n"
                                        text += f"  Std Dev: {values.std():.2f}\n"
                                        text += f"  Min: {values.min():.2f}\n"
                                        text += f"  Max: {values.max():.2f}\n"
                            
                            text_cols = df.select_dtypes(include=['object']).columns
                            if len(text_cols) > 0:
                                text += "\nTEXT COLUMNS:\n"
                                text += "-" * 80 + "\n"
                                for col in text_cols:
                                    unique_count = df[col].nunique()
                                    text += f"  {col}: {unique_count} unique values\n"
                    
                    except Exception as stats_error:
                        text += f"\n(Could not compute statistics: {str(stats_error)})\n"
                    
                    text += "\n"
                
                wb.close()
                
            else:
                text = DocumentProcessor._extract_excel_with_pandas(file_path, display_name)
            
        except Exception as e:
            st.error(f"Error reading Excel file: {str(e)}")
            text = f"[Error reading Excel file: {str(e)}]\n"
            try:
                text += "\nAttempting fallback extraction...\n"
                display_name = original_filename if original_filename else Path(file_path).name
                text += DocumentProcessor._extract_excel_with_pandas(file_path, display_name)
            except:
                pass
        
        return text
    
    @staticmethod
    def _extract_excel_with_pandas(file_path: str, original_filename: str = None) -> str:
        """Fallback method to extract Excel using pandas only"""
        text = ""
        try:
            display_name = original_filename if original_filename else Path(file_path).name
            
            file_extension = Path(file_path).suffix.lower()
            engine = 'openpyxl' if file_extension == '.xlsx' else 'xlrd'
            
            xls = pd.ExcelFile(file_path, engine=engine)
            
            text += f"\n{'='*80}\n"
            text += f"DOCUMENT: {display_name}\n"
            text += f"{'='*80}\n"
            text += f"Excel File (Pandas Extraction)\n\n"
            
            for sheet_name in xls.sheet_names:
                try:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    
                    if df.empty or df.dropna(how='all').empty:
                        continue
                    
                    text += f"\n{'='*80}\n"
                    text += f"SHEET: {sheet_name}\n"
                    text += f"{'='*80}\n\n"
                    
                    text += f"Columns: {list(df.columns)}\n"
                    text += f"Rows: {len(df)}\n\n"
                    
                    text += "DATA:\n"
                    text += "-" * 80 + "\n"
                    text += df.fillna('').to_string(index=False) + "\n\n"
                    
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        text += "STATISTICS:\n"
                        text += "-" * 80 + "\n"
                        for col in numeric_cols:
                            values = df[col].dropna()
                            if len(values) > 0:
                                text += f"{col}: Sum={values.sum():.2f}, Mean={values.mean():.2f}, Min={values.min():.2f}, Max={values.max():.2f}\n"
                        text += "\n"
                    
                except Exception as sheet_error:
                    text += f"\n[Sheet '{sheet_name}' error: {str(sheet_error)}]\n\n"
            
            xls.close()
            
        except Exception as e:
            text += f"[Pandas extraction error: {str(e)}]"
        
        return text
    
    @staticmethod
    def process_document(uploaded_file) -> Dict:
        """Process uploaded document and extract text"""
        original_filename = uploaded_file.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(original_filename).suffix) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        text = ""
        file_extension = Path(original_filename).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                text = DocumentProcessor.extract_text_from_pdf(tmp_path, original_filename)
            elif file_extension in ['.docx', '.doc']:
                text = DocumentProcessor.extract_text_from_docx(tmp_path, original_filename)
            elif file_extension in ['.xlsx', '.xls']:
                text = DocumentProcessor.extract_text_from_excel(tmp_path, original_filename)
            else:
                text = ""
                st.warning(f"Unsupported file type: {file_extension}")
        except Exception as e:
            st.error(f"Error processing {original_filename}: {str(e)}")
            text = ""
        finally:
            try:
                os.unlink(tmp_path)
            except PermissionError:
                time.sleep(0.1)
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
        
        return {
            'filename': original_filename,
            'text': text,
            'file_type': file_extension
        }

# ===================================================================
# RAG SYSTEM
# ===================================================================

class RAGSystem:
    """RAG system for document retrieval and question answering"""
    
    def __init__(self, api_key: str):
        """Initialize RAG system with Gemini API"""
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-3-flash-preview')
        self.vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
    
    def create_vector_store(self, documents: List[Dict]):
        """Create vector store from documents"""
        self.documents = documents
        self.document_texts = [doc['text'] for doc in documents]
        self.document_names = [doc['filename'] for doc in documents]
        
        if self.document_texts:
            self.document_vectors = self.vectorizer.fit_transform(self.document_texts)
            return True
        return False
    
    def retrieve_relevant_documents(self, query: str, top_k: int = 7) -> List[Dict]:
        """Retrieve most relevant documents for the query"""
        query_vector = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vector, self.document_vectors).flatten()
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        relevant_docs = []
        for idx in top_indices:
            if similarities[idx] > 0.03:
                relevant_docs.append({
                    'filename': self.document_names[idx],
                    'text': self.document_texts[idx],
                    'similarity': similarities[idx]
                })
        
        return relevant_docs
    
    def generate_answer(self, query: str, context_docs: List[Dict]) -> str:
        """Generate answer using Gemini API with retrieved context"""
        context = ""
        has_excel = False
        excel_files = []
        
        for doc in context_docs:
            context += f"\n{'='*80}\n"
            context += f"Document: {doc['filename']}\n"
            context += f"Relevance Score: {doc['similarity']:.2%}\n"
            context += f"{'='*80}\n"
            
            if '.xlsx' in doc['filename'].lower() or '.xls' in doc['filename'].lower():
                context += doc['text'][:100000] + "\n"
                has_excel = True
                excel_files.append(doc['filename'])
            else:
                context += doc['text'][:80000] + "\n"
        
        formula_keywords = ['formula', 'calculate', 'calculation', 'compute', 'excel', 'cell']
        is_formula_query = any(keyword in query.lower() for keyword in formula_keywords)
        
        threat_keywords = ['threat', 'vulnerability', 'vulnerabilities', 'risk', 'annexure', 'list']
        is_threat_query = any(keyword in query.lower() for keyword in threat_keywords)

        excel_instruction = ""
        if has_excel:
            excel_priority = ""
            if is_formula_query:
                excel_priority = f"""
**🔴 CRITICAL - FORMULA QUERY DETECTED:**
This question asks about formulas or calculations.
Excel files in context: {', '.join(excel_files)}
PRIORITY: Look in Excel files FIRST before DOCX files!

When answering about formulas:
1. FIRST check "FORMULAS FOUND IN THIS SHEET" sections in Excel files
2. Extract actual Excel formulas (Cell X: =...)
3. Show formula AND result
4. ONLY if no Excel formula found, then check DOCX explanations

DO NOT give conceptual formulas from DOCX when actual Excel formulas exist!
"""
            
            excel_instruction = excel_priority + """
**EXCEL FILE HANDLING:**
- Look for "FORMULAS FOUND IN THIS SHEET" sections
- Extract formulas that start with "Cell" and "="
- Include calculation formulas like =A1*B1, =VLOOKUP(...), =SUM(), =IF(...), etc.
- Show both the formula AND the result
- Specify cell coordinates (e.g., Cell F2, Cell G5)
- If question asks "how to calculate", MUST show actual Excel formula first
"""

        threat_instruction = ""
        if is_threat_query:
            threat_instruction = """
**🔴 THREAT/VULNERABILITY QUERY DETECTED:**
This question asks about threats and vulnerabilities.

CRITICAL INSTRUCTIONS:
1. Look for "Threat Vulnerability database" sheet in Excel files
2. Look for "ANNEXURE A" or "Annexure A" sections in DOCX files  
3. Look for tables containing threat and vulnerability data
4. Extract ALL rows from these sections
5. List each threat with its associated vulnerability
6. DO NOT say "content not provided" if you see the sheet name or section title - the data IS there!
"""
   
        prompt = f"""You are an expert assistant specialized in Risk Management, Asset Management, and Information Security standards (ISO 27005, ISO 31000, NIST). 

Your task is to answer questions based ONLY on the provided documents with maximum detail and accuracy.

CONTEXT FROM DOCUMENTS:
{context}

QUESTION: {query}

{excel_instruction}

{threat_instruction}

CRITICAL INSTRUCTIONS:
1. **PRIORITIZE EXCEL FORMULAS**: If question asks about formulas/calculations and Excel files are in context:
   - Look in Excel files FIRST (search for "FORMULAS FOUND" sections)
   - Extract actual formulas: Cell F2: =D2*E2 → Result: 20
   - Show cell references, formula, and calculated value
   - ONLY after showing Excel formulas, you can add DOCX explanations

2. **BE SPECIFIC AND DETAILED**: Never say "no specific details provided" or "formula not provided" - if it exists in any document, extract it completely

3. **INCLUDE ALL EXAMPLES**: If examples, tables, or data are present, include them

4. **USE EXACT TERMINOLOGY**: Use exact terms from documents

5. **CITE ACCURATELY**: 
   - For Excel: "From Risk_Assessment_Template.xlsx, Sheet 'Risk Assessment', Cell F2: =D2*E2"
   - For DOCX: "From Risk_Management_Guidelines.docx, Section 4.5.1"

6. **BE COMPREHENSIVE**: Provide complete information, not summaries

7. **STRUCTURE CLEARLY**: Use formatting for clarity

8. **NO GENERIC ANSWERS**: Extract ALL relevant information

If the answer truly cannot be found after thoroughly checking all documents including Excel formulas, state: "This specific information is not found in the provided documents."

ANSWER (prioritize Excel formulas for calculation questions):"""
        
        try:
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error generating answer: {str(e)}"

# ===================================================================
# MAIN APP
# ===================================================================

def main():
    st.title("📚 Persistent RAG Document Q&A System")
    st.markdown("Upload your documents **ONCE** and use forever! - powered by Google Gemini")
    
    # Check if knowledge base exists on disk on first load
    if not st.session_state.kb_loaded:
        if knowledge_base_exists():
            with st.spinner("🔄 Loading saved knowledge base..."):
                documents, vectorizer, document_vectors, metadata = load_knowledge_base()
                
                if documents is not None:
                    st.session_state.documents = documents
                    st.session_state.vectorizer = vectorizer
                    st.session_state.document_vectors = document_vectors
                    st.session_state.processed = True
                    st.session_state.kb_loaded = True
                    
                    # Show success message
                    st.success(f"✅ Knowledge base loaded! {metadata['num_documents']} documents ready")
                    st.info(f"📅 Last saved: {metadata['saved_at']}")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key input
        api_key = st.text_input(
            "Enter Google Gemini API Key",
            type="password",
            help="Get your API key from https://makersuite.google.com/app/apikey"
        )
        
        st.markdown("---")
        
        # Show knowledge base status
        if st.session_state.processed:
            st.success("✅ Knowledge Base: READY")
            st.info(f"📄 {len(st.session_state.documents)} documents loaded")
            
            # Show document list
            with st.expander("📋 Loaded Documents"):
                for doc in st.session_state.documents:
                    st.markdown(f"• {doc['filename']}")
            
            st.markdown("---")
            
            # Delete knowledge base button
            if st.button("🗑️ Delete Knowledge Base", type="secondary"):
                if delete_knowledge_base():
                    st.session_state.documents = []
                    st.session_state.processed = False
                    st.session_state.chat_history = []
                    st.session_state.vectorizer = None
                    st.session_state.document_vectors = None
                    st.session_state.kb_loaded = False
                    st.success("✅ Knowledge base deleted!")
                    st.rerun()
        
        else:
            st.warning("⚠️ No knowledge base found")
            st.markdown("Upload and process documents below")
        
        st.markdown("---")
        
        st.header("📤 Upload Documents")
        st.markdown("Upload documents to create/update knowledge base")
        
        uploaded_files = st.file_uploader(
            "Choose files",
            type=['pdf', 'docx', 'doc', 'xlsx', 'xls'],
            accept_multiple_files=True,
            help="Upload your documents here"
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} file(s) selected")
            
            if st.button("🔄 Process & Save Documents", type="primary"):
                if not api_key:
                    st.error("Please enter your Gemini API key first!")
                else:
                    with st.spinner("Processing documents..."):
                        # Process all documents
                        documents = []
                        progress_bar = st.progress(0)
                        
                        for idx, uploaded_file in enumerate(uploaded_files):
                            st.info(f"Processing: {uploaded_file.name}")
                            doc_data = DocumentProcessor.process_document(uploaded_file)
                            documents.append(doc_data)
                            progress_bar.progress((idx + 1) / len(uploaded_files))
                        
                        # Create RAG system
                        rag_system = RAGSystem(api_key)
                        
                        # Create vector store
                        success = rag_system.create_vector_store(documents)
                        
                        if success:
                            # Save to disk
                            if save_knowledge_base(documents, rag_system.vectorizer, rag_system.document_vectors):
                                st.session_state.documents = documents
                                st.session_state.vectorizer = rag_system.vectorizer
                                st.session_state.document_vectors = rag_system.document_vectors
                                st.session_state.processed = True
                                st.session_state.kb_loaded = True
                                
                                st.success("✅ Documents processed and saved to disk!")
                                st.success("💾 Knowledge base will persist across sessions!")
                                st.rerun()
                            else:
                                st.error("❌ Failed to save knowledge base")
                        else:
                            st.error("Failed to process documents")
    
    # Main content area
    if not st.session_state.processed:
        st.info("👈 Please upload documents and process them using the sidebar")
        
        # Display instructions
        with st.expander("📖 How to use this system"):
            st.markdown("""
            ### Steps to get started:
            
            1. **Enter API Key**: Get your free API key from [Google AI Studio](https://makersuite.google.com/app/apikey)
            2. **Upload Documents**: Upload your documents (PDF, DOCX, or Excel files)
            3. **Process & Save**: Click "Process & Save Documents" button
            4. **Use Forever**: Knowledge base saved to disk - no need to upload again!
            
            ### Features:
            - ✅ **Persistent Storage**: Upload once, use forever
            - ✅ Supports PDF, DOCX, and Excel files
            - ✅ Intelligent document retrieval
            - ✅ Accurate answers based on your documents
            - ✅ Source citations
            - ✅ Chat history
            
            ### Knowledge Base Status:
            - 📁 Saved in: `knowledge_base/` folder
            - 💾 Includes: Documents, vectors, and metadata
            - 🔄 Auto-loads on app restart
            """)
    
    else:
        # Display processed documents
        with st.expander("📑 View Loaded Documents"):
            for doc in st.session_state.documents:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**{doc['filename']}**")
                with col2:
                    st.markdown(f"`{doc['file_type']}`")
                
                preview = doc['text'][:200] + "..." if len(doc['text']) > 200 else doc['text']
                st.text(preview)
                st.markdown("---")
        
        # Chat interface
        st.header("💬 Ask Questions")
        
        if not api_key:
            st.warning("⚠️ Please enter your Gemini API key in the sidebar to ask questions")
        else:
            # Initialize RAG system for this session
            if 'rag_system' not in st.session_state:
                rag_system = RAGSystem(api_key)
                rag_system.documents = st.session_state.documents
                rag_system.document_texts = [doc['text'] for doc in st.session_state.documents]
                rag_system.document_names = [doc['filename'] for doc in st.session_state.documents]
                rag_system.vectorizer = st.session_state.vectorizer
                rag_system.document_vectors = st.session_state.document_vectors
                st.session_state.rag_system = rag_system
            
            # Display chat history
            for chat in st.session_state.chat_history:
                with st.chat_message("user"):
                    st.write(chat['question'])
                with st.chat_message("assistant"):
                    st.write(chat['answer'])
                    if 'sources' in chat:
                        with st.expander("📎 Sources"):
                            for source in chat['sources']:
                                st.markdown(f"- **{source['filename']}** (Relevance: {source['similarity']:.2%})")
            
            # Question input
            question = st.chat_input("Ask a question about your documents...")
            
            if question:
                with st.chat_message("user"):
                    st.write(question)
                
                with st.chat_message("assistant"):
                    with st.spinner("Searching documents and generating answer..."):
                        relevant_docs = st.session_state.rag_system.retrieve_relevant_documents(question, top_k=7)
                        
                        if not relevant_docs:
                            answer = "I couldn't find relevant information in the uploaded documents to answer this question."
                            st.write(answer)
                        else:
                            answer = st.session_state.rag_system.generate_answer(question, relevant_docs)
                            st.write(answer)
                            
                            with st.expander("📎 Sources"):
                                for doc in relevant_docs:
                                    st.markdown(f"- **{doc['filename']}** (Relevance: {doc['similarity']:.2%})")
                            
                            st.session_state.chat_history.append({
                                'question': question,
                                'answer': answer,
                                'sources': relevant_docs
                            })

if __name__ == "__main__":
    main()